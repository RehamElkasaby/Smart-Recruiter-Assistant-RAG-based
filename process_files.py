# process_files.py
import os
from typing import List, Tuple
import pdfplumber
from docx import Document
from langchain.schema import Document as LangDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

def extract_name_from_filename(filename: str) -> str:
    base = os.path.splitext(filename)[0]
    # Remove special characters and normalize
    return ' '.join(word.capitalize() for word in base.replace('_', ' ').replace('-', ' ').split())

def read_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text(x_tolerance=1, y_tolerance=1) + "\n"
    return text

def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)

def load_documents(file_path: str) -> LangDocument:
    if file_path.endswith('.pdf'):
        text = read_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = read_docx(file_path)
    else:
        with open(file_path, 'r') as f:
            text = f.read()

    if not text:
        raise ValueError(f"Empty content in {file_path}")

    filename = os.path.basename(file_path)
    candidate_name = extract_name_from_filename(filename)

    return LangDocument(
        page_content=text,
        metadata={
            "source": filename,
            "candidate_name": candidate_name,  # Store extracted name
            "document_type": "full_cv"
        }
    )

def split_documents(doc: LangDocument) -> Tuple[List[LangDocument], str]:
    if not doc:
        return [], ""

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    splits = text_splitter.split_documents([doc])

    cv_name = doc.metadata["source"]
    candidate_name = doc.metadata["candidate_name"]

    for split in splits:
        split.metadata.update({
            "is_section": True,
            "full_doc_reference": cv_name,
            "candidate_name": candidate_name,
            "document_type": "cv_section"
        })
    return splits, candidate_name

def process_uploaded_files(file_paths: List[str]) -> Tuple[List[LangDocument], List[str]]:
    all_chunks = []
    all_candidate_names = []

    for file_path in file_paths:
        full_doc = load_documents(file_path)
        chunks, candidate_name = split_documents(full_doc)
        all_chunks.extend(chunks)
        all_candidate_names.append(candidate_name)
    return all_chunks, all_candidate_names